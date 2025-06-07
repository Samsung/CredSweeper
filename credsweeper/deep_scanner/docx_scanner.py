import io
import logging
from abc import ABC
from typing import List, Optional

import docx
from docx.document import Document
from docx.oxml import CT_P, CT_Tbl, CT_SectPr, CT_TcPr
from docx.section import Section, _Header, _Footer
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from lxml.etree import _Element

from credsweeper.credentials.candidate import Candidate
from credsweeper.deep_scanner.abstract_scanner import AbstractScanner
from credsweeper.file_handler.data_content_provider import DataContentProvider
from credsweeper.file_handler.string_content_provider import StringContentProvider

logger = logging.getLogger(__name__)


class DocxScanner(AbstractScanner, ABC):
    """Implements docx scanning"""

    @staticmethod
    def _iter_block_items(block):
        if isinstance(block, Paragraph):
            yield block
            return
        if isinstance(block, (_Header, _Footer)):
            for table in block.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from DocxScanner._iter_block_items(cell)
            for paragraph in block.paragraphs:
                yield paragraph
            return
        elif isinstance(block, Document):
            parent_elm = block.element.body
        elif isinstance(block, Section):
            yield from DocxScanner._iter_block_items(block.header)
            yield from DocxScanner._iter_block_items(block.footer)
            return
        elif isinstance(block, _Cell):
            parent_elm = block._tc  # pylint: disable=W0212
        else:
            raise ValueError(f"unrecognised:{type(block)}")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, block)
            elif isinstance(child, CT_Tbl):
                table = Table(child, block)
                for row in table.rows:
                    for cell in row.cells:
                        yield from DocxScanner._iter_block_items(cell)
            elif isinstance(child, (CT_TcPr, CT_SectPr)):
                # config
                pass
            elif isinstance(child, _Element):
                yield child
            else:
                logger.warning(f"Unknown:{type(child)}")

    def data_scan(
            self,  #
            data_provider: DataContentProvider,  #
            depth: int,  #
            recursive_limit_size: int) -> Optional[List[Candidate]]:
        """Tries to scan DOCX text with splitting by lines"""
        try:
            docx_lines: List[str] = []

            doc = docx.Document(io.BytesIO(data_provider.data))
            for block in self._iter_block_items(doc):
                if block.text:
                    docx_lines.append(block.text)

            header_lines_set = set()
            footer_lines_set = set()
            for section in doc.sections:
                for header in [section.first_page_header, section.even_page_header, section.header]:
                    for block in self._iter_block_items(header):
                        if block.text:
                            header_lines_set.add(block.text)
                for footer in [section.first_page_footer, section.even_page_footer, section.footer]:
                    for block in self._iter_block_items(footer):
                        if block.text:
                            footer_lines_set.add(block.text)
            docx_lines.extend(sorted(list(header_lines_set)))
            docx_lines.extend(sorted(list(footer_lines_set)))

            string_data_provider = StringContentProvider(lines=docx_lines,
                                                         file_path=data_provider.file_path,
                                                         file_type=data_provider.file_type,
                                                         info=f"{data_provider.info}|DOCX")
            docx_candidates = self.scanner.scan(string_data_provider)
            return docx_candidates

        except Exception as docx_exc:
            logger.error(f"{data_provider.file_path}:{docx_exc}")
        return None
